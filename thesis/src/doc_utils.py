#!pip install python-docx
import docx
import os, sys
import glob
import re
import string
import defines

import numpy as np
import pandas as pd
import seaborn as sns
from nltk import tokenize


pd.options.display.float_format = '{:f}'.format

doc_cols = ['path','file_name','client_tag','therapist_tag','num_par']
doc_db = pd.DataFrame(columns=doc_cols)
par_db =  pd.DataFrame()
plane_par_db = pd.DataFrame()
block_db = pd.DataFrame()
sent_db = pd.DataFrame()
# utils for files 

# def check_par_types_order():
#     error_order= {
#         'middle': ['whole', 'non_nar'],
#         'start' : ['whole', 'non_nar'],
#         'end' : ['middle']
#         'non_nar' : ['whole','middle,'end']
#         }  




def get_random_paragraph(query):
    match =  plane_par_db.query(query)
    return match.sample()

def count_narr_per_par(par,is_nar):
    return max(par.count(defines.START_CHAR),par.count(defines.END_CHAR),is_nar)

def check_text_for_illegal_labels(doc_idx,par):
        # capture if there is [start:start] or [end:end]
    error_pattern_st = "(" + defines.START_CHAR + "(?:(?!" + defines.END_CHAR + ").)*" +  defines.START_CHAR +")";
    error_pattern_end = "(" + defines.END_CHAR + "(?:(?!" + defines.START_CHAR + ").)*" +  defines.END_CHAR +")";
    error_pattern = "(" + error_pattern_st + "|" +  error_pattern_end + ")";
    if (re.search(error_pattern,par)):
        print("Doc {}\nstring gave ERROR\n{}".format(doc_idx,par))
        return 1
    return 0

def split_block_to_sentences(text):
    sent_list = tokenize.sent_tokenize(text)
    for i,item in enumerate(sent_list):
        clean_item = clean_text(item)
        if len(clean_item) != 0: # disregard empty strings
            sent_list[i] = clean_item
    return sent_list

def add_sentences_of_blocks_to_db(block_db_idx):
    global sent_db
    block_line = block_db.iloc[block_db_idx]
    block = block_line['text']
    sent_list = split_block_to_sentences(block)
    for i,sentence in enumerate(sent_list):
        curr_db_idx = sent_db.shape[0]
        sent_db.loc[curr_db_idx,'text'] = sentence
        sent_db.loc[curr_db_idx,'sent_idx_in_block'] = i
        sent_db.loc[curr_db_idx,'block_idx'] = block_db_idx
        sent_db.loc[curr_db_idx,'is_nar'] = block_line['is_nar']
        sent_db.loc[curr_db_idx,'doc_idx'] = block_line['doc_idx']
        sent_db.loc[curr_db_idx,'par_db_idx'] = block_line['par_db_idx']
        sent_db.loc[curr_db_idx,'par_idx_in_doc'] = block_line['par_idx_in_doc']
        sent_db.loc[curr_db_idx,'par_type'] = block_line['par_type']
        sent_db.loc[curr_db_idx,'block_type'] = block_line['block_type']
        sent_db.loc[curr_db_idx,'nar_idx'] = block_line['nar_idx']
        sent_db.loc[curr_db_idx,'sent_len'] = len(sentence)

        

def split_par_to_blocks_keep_order(plane_par_db_idx):
    par = plane_par_db.loc[plane_par_db_idx,'text']
    startNum = par.count(defines.START_CHAR)
    endNum = par.count(defines.END_CHAR)
    block_list = []
    tag = ""
    outside_nar = ""
    splited = []

    if startNum == 0 and endNum == 0: # text is missing start and end symbols
        if plane_par_db.loc[plane_par_db_idx,'is_nar'] == 0: #entire paragraph is not narrative
            tag = "not_nar"
        else: #entire paragraph is narrative
            tag = "middle"
        # par = clean_text(par)
        block_list.insert(0,(tag,par))
    else:
        splited = re.split('(&|#)', par) # used for keeping original order between blocks
        splited_clean = splited
        for i,block in enumerate(splited):
            if '%' in block: # TBD handle story summary
                continue
            splited_clean[i] = clean_text(block)
        my_regex = {
            'whole' :defines.START_CHAR + ".*?" + defines.END_CHAR,     # [start:end] 
            'start' : defines.START_CHAR + ".*", # [start:]
            'end' :".*" +  defines.END_CHAR # [:end]
        }
        outside_nar = par
        for tag,regex in my_regex.items():
            nar_blocks = re.findall(regex,outside_nar)
            for j,block in enumerate(nar_blocks):
                if len(block) !=0:
                    # block = clean_text(block)
                    block_idx = get_index_of_block_in_par(splited_clean,block,plane_par_db_idx)
                    splited[block_idx] = "" # erase narrative blocks from splited paragraph
                    block_list.insert(block_idx,(tag,block))
            outside_nar = re.sub(regex,'',outside_nar)
        
        # handle the rest items in list - that must be non-narrative
        for i,block in enumerate(splited):
            if len(block)!=0:
        # if len(outside_nar) !=0 :
                if '%' in block:
                    continue # TBD handle story summary
                # block = clean_text(block)
                block_idx = get_index_of_block_in_par(splited_clean,block,plane_par_db_idx)
                block_list.insert(block_idx,("not_nar",block))
    return block_list

def get_index_of_block_in_par(splited,block,plane_par_db_idx):
    block = clean_text(block)
    if not block in splited:
        print("{} \n par[{}]not in \n{}".format(plane_par_db_idx,block,splited))
        return -1
    else:
        return splited.index(block)
    
def split_par_to_blocks(plane_par_db_idx):
    par = plane_par_db.loc[plane_par_db_idx,'text']
    startNum = par.count(defines.START_CHAR)
    endNum = par.count(defines.END_CHAR)
    block_list = []
    tag = ""
    outside_nar = ""
    if startNum == 0 and endNum == 0: # text is missing start and end symbols
        if plane_par_db.loc[plane_par_db_idx,'is_nar'] == 0: #entire paragraph is not narrative
            tag = "not_nar"
        else: #entire paragraph is narrative
            tag = "middle"
        par = clean_text(par)
        block_list.append((tag,par))
    else:
        my_regex = {
            'whole' :defines.START_CHAR + ".*?" + defines.END_CHAR,     # [start:end] 
            'start' : defines.START_CHAR + ".*", # [start:]
            'end' :".*" +  defines.END_CHAR # [:end]
        }
        outside_nar = par
        for tag,regex in my_regex.items():
            nar_blocks = re.findall(regex,outside_nar)
            for j,block in enumerate(nar_blocks):
                if len(block) !=0:
                    block = clean_text(block)
                    block_list.append((tag,block))
            outside_nar = re.sub(regex,'',outside_nar)
    if len(outside_nar) !=0 :
        outside_nar= clean_text(outside_nar)
        block_list.append(("not_nar",outside_nar))
    return block_list

def add_features_prev_is_nar():
    global block_db
    block_db['one_before_is_nar']=block_db['is_nar'].shift(periods=1, fill_value=0)
    block_db['two_before_is_nar']=block_db['is_nar'].shift(periods=2, fill_value=0)


def get_last_nar_idx_from_block_db():
    global block_db
    return block_db['nar_idx'].max()

def add_blocks_of_par_to_db(plane_par_db_idx):
    global plane_par_db, block_db
    is_nar = 0
    # block_list = split_par_to_blocks(plane_par_db_idx)
    block_list = split_par_to_blocks_keep_order(plane_par_db_idx)
    par_db_line = plane_par_db.iloc[plane_par_db_idx]
    for i,tupple in enumerate(block_list):
        curr_db_idx = block_db.shape[0]
        curr_nar_idx = 0 if curr_db_idx == 0 else get_last_nar_idx_from_block_db() 
        if tupple[0] in ['start','whole']:
            curr_nar_idx+=1
        is_nar = 1 if tupple[0] != 'not_nar' else 0
        block_db.loc[curr_db_idx,'text'] = tupple[1]
        block_db.loc[curr_db_idx,'is_nar'] = is_nar
        block_db.loc[curr_db_idx,'doc_idx'] = par_db_line['doc_idx']
        block_db.loc[curr_db_idx,'par_idx_in_doc'] = par_db_line['par_idx_in_doc']
        block_db.loc[curr_db_idx,'par_db_idx'] = plane_par_db_idx
        block_db.loc[curr_db_idx,'par_type'] = par_db_line['par_type']
        block_db.loc[curr_db_idx,'block_type'] = tupple[0]
        block_db.loc[curr_db_idx,'nar_idx'] = curr_nar_idx if is_nar else 0


def save_all_blocks():
    global plane_par_db,block_db
    for i in plane_par_db.index:
        add_blocks_of_par_to_db(i)
    block_db.to_csv("block_db.csv",index=False)
    print("All blocks saved")

def save_all_sentences():
    global block_db, sent_db
    for i in block_db.index:
        add_sentences_of_blocks_to_db(i)
    sent_db.to_csv("sent_db.csv",index=False)
    print("All sentences saved")

def split_doc_to_paragraphs(doc,doc_idx):
    global par_db, plane_par_db
    inside_narrative = 0
    nar_idx = 0
    for i,par in enumerate(doc.paragraphs):
        curr_par_db_idx = plane_par_db.shape[0]
        text,par_type = get_par_type_erase(par.text,doc_idx)
        if par_type == 'empty':
#             print("{} is empty: {}".format(i,par.text))
            continue
        plane_par_db.loc[curr_par_db_idx,'doc_idx'] = doc_idx
        plane_par_db.loc[curr_par_db_idx,'text'] = text
        plane_par_db.loc[curr_par_db_idx,'par_len'] = len(text)
        if par_type == 'no_mark':
            par_type = check_unknown_par_type(curr_par_db_idx)
        plane_par_db.loc[curr_par_db_idx,'par_type'] = par_type
        plane_par_db.loc[curr_par_db_idx,'par_idx_in_doc'] = i
        if defines.START_CHAR in par.text:
            inside_narrative = 1
            nar_idx+=1 # starting narrative indexing from 1
        plane_par_db.loc[curr_par_db_idx,'is_nar'] = inside_narrative
        plane_par_db.loc[curr_par_db_idx,'nar_per_par'] = count_narr_per_par(text,inside_narrative)
        plane_par_db.loc[curr_par_db_idx,'nar_idx'] = nar_idx if inside_narrative else 0
        if defines.END_CHAR in par.text and not (defines.START_CHAR in par.text):
            inside_narrative = 0
    print("Doc {} paragraphs saved".format(doc_idx))

def check_unknown_par_type(curr_par_db_idx):
    one_before_idx = curr_par_db_idx - 1
    two_before_idx = curr_par_db_idx - 2
    if one_before_idx in plane_par_db.index and plane_par_db.loc[one_before_idx,'par_type']=='segment':
        par_type = plane_par_db.loc[two_before_idx,'par_type']
    else:
        par_type = 'no_mark'
    return par_type

def save_all_docs_paragraphs():
    global doc_db
    for doc_idx in doc_db.index:
        doc = docx.Document(doc_db.loc[doc_idx,'path'])
        split_doc_to_paragraphs(doc,doc_idx)




def save_docs_db():
    global doc_db
    doc_path_list = get_labeled_files()
    doc_obj_list = get_doc_objects(doc_path_list)
    for path in doc_path_list:
        add_doc_to_db(path)
    doc_db.to_csv("doc_db.csv",index=False)

def remove_punctuation(_text):
    text = _text.translate(str.maketrans('', '',string.punctuation))
    return text
 

def get_labeled_files():
    doc_path_list = []
    for file in glob.glob("./tmp/*_l.docx"): # _l is name pattern of labeled *.docx files
        doc_path_list.append(file)
    return doc_path_list

def get_doc_objects(doc_path_list):
    doc_list = []
    for path in doc_path_list: 
        doc_list.append(docx.Document(path))
    return doc_list

def add_doc_to_db(path):
    global doc_db
    file_name = os.path.basename(path)
    doc = docx.Document(path)
    client_tag, therapist_tag = get_client_therapist_tag(doc)
    num_par = len(doc.paragraphs)
    doc_list = [path,file_name,client_tag,therapist_tag,num_par]
    doc_db.loc[doc_db.shape[0]] = doc_list


def get_client_therapist_tag(doc):
    client_tag = ''
    therapist_tag = ''
    for par in doc.paragraphs[:4]:
        if 'CLIENT' in par.text:
            client_tag = par.text.split()[0]
        if 'THERAPIST' in par.text:
            therapist_tag = par.text.split()[0]
    return client_tag,therapist_tag

def add_length_of_paragraphs(_df):
    df = _df.copy()
    df['par_len'] = df['text'].str.len()
    return df

def add_length_of_nar_in_words(_df):
    _df.loc[_df['nar_len'] is not None,['nar_len_words']] = _df.groupby('nar_idx')['par_len'].sum()

def clean_text(text):
    text= text.replace(r'[@,#,&,\*,\t]*','')
    text=text.replace('\t','')
    text=remove_punctuation(text)
    return text

def get_par_type_erase(par,doc_idx,do_clean=False):
    client_tag = doc_db.loc[doc_idx,'client_tag']
    therapist_tag = doc_db.loc[doc_idx,'therapist_tag']
    segment_string = "(סגמנט|דקה)" + ".*[0-9]"
    par_type = 'no_mark'
    if len(par) == 0:
        par_type = 'empty'
    if client_tag in par[:20]: # search for a tag in the begginning of a line
        par = par.replace(client_tag, '')
        par_type = 'client'
    if therapist_tag in par[:20]: # search for a tag in the begginning of a line
        par = par.replace(therapist_tag, '')
        par_type= 'therapist'
    if re.search(segment_string,par[:20]):
        par_type ='segment'
    if not re.search(r'[א-ת]',par) and re.search(r'\d',par): # if there is no hebrew letters and there is numbers
        par_type = 'segment'
#     if '%' in par:
#         par = par.replace('%', '')
#         par_type= 'summary' # TBD implement summary extraction
    if 'CLIENT' in par or 'THERAPIST' in par:
        par_type = 'no_mark'
    check_text_for_illegal_labels(doc_idx,par)
    if(do_clean):
        par = clean_text(par)
    return par,par_type



# def add_paragraphs_to_db(doc_idx,doc_db,par_db,sent_db):
#     doc = docx.Document(doc_db.loc[doc_idx,'path'])
#     inside_narrative  = 0
#     narrative_idx = -1 #index of narrative within given doc
#     idx_in_nar = -1 #index of paragraph within given narrative
#     par_idx = -1 # index of paragraph within given doc
#     pre_par_speaker = ''
#     glob_nar_index = ''
#     for i,par in enumerate(doc.paragraphs):
#         curr_par_db_idx = par_db.shape[0]
#         text,par_type = get_par_type_erase(par.text,doc_idx,doc_db)
#         if par_type == 'summary':
#             continue; # TBD save summary in different db for future use
#         par_db.loc[curr_par_db_idx,'doc_idx'] = doc_idx
#         par_db.loc[curr_par_db_idx,'text'] = text
#         par_db.loc[curr_par_db_idx,'par_len'] = len(text)
#         par_db.loc[curr_par_db_idx,'par_type'] = par_type
#         par_idx+=1
#         par_db.loc[curr_par_db_idx,'par_idx'] = par_idx
#         if par_type == 'no_mark':
#             if curr_par_db_idx-1 in par_db.index and par_db.loc[curr_par_db_idx-1,'par_type']=='segment':
#                 par_db.loc[curr_par_db_idx,'par_type'] = par_db.loc[curr_par_db_idx-2,'par_type']       
#         if par_db.loc[curr_par_db_idx,'par_type'] in ['segment','no_mark']:
#             par_db.loc[curr_par_db_idx,'is_nar'] = 0
#             print("skipping for doc {} entry {} type {} text {}".format(doc_idx,curr_par_db_idx,par_db.loc[curr_par_db_idx,'par_type'],text))
#             continue # skip parahraph processing
#         if defines.START_CHAR in par.text:
#             inside_narrative = 1
#             idx_in_nar = 0
#             narrative_idx+=1
#         par_db.loc[curr_par_db_idx,'is_nar'] = inside_narrative
#         par_db.loc[curr_par_db_idx,'nar_idx'] = narrative_idx if (inside_narrative) else None
#         glob_nar_index = f"{doc_idx}_{narrative_idx}" if (inside_narrative) else None
#         par_db.loc[curr_par_db_idx,'glob_nar_idx'] = glob_nar_index
#         par_db.loc[curr_par_db_idx,'idx_in_nar'] = idx_in_nar if (inside_narrative) else None
#         idx_in_nar+=1
#         if defines.END_CHAR in par.text:
#             print ("update nar {} len to {}".format(narrative_idx,idx_in_nar))
#             par_db.loc[par_db['glob_nar_idx']==glob_nar_index,'nar_len'] = idx_in_nar
#             par_db.loc[par_db['glob_nar_idx']==glob_nar_index,'nar_len_words'] = par_db.loc[par_db['glob_nar_idx']==glob_nar_index,'par_len'].sum()
#             inside_narrative = 0
            
# def add_sentence_to_db(sent_list,doc_idx,par_idx,is_nar,glob_nar_index,par_type,sent_db):
#     for i,sent in enumerate(sent_list):
#         curr_idx = sent_db.shape[0]
#         sent= clean_text(sent)
#         sent_db.loc[curr_idx,'text']=sent
#         sent_db.loc[curr_idx,'is_nar']=is_nar
#         sent_db.loc[curr_idx,'glob_nar_index']=glob_nar_index
#         sent_db.loc[curr_idx,'par_idx']=par_idx
#         sent_db.loc[curr_idx,'sent_len']=len(sent)
#         sent_db.loc[curr_idx,'par_type']=par_type
#         sent_db.loc[curr_idx,'doc_idx']=doc_idx
    
    


