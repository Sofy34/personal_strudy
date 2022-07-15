#!pip install python-docx
import docx
import os
import sys
import glob
import re
import string
import defines

import numpy as np
import pandas as pd
import seaborn as sns
from nltk import tokenize
import common_utils

pd.options.display.float_format = "{:f}".format

global doc_db
global par_db
global block_db
global sent_db
global debug_db


def get_random_paragraph(query):
    match = par_db.query(query)
    return match.sample()


def text_contains_char(text):
    has_char = re.search(r"[a-z,A-Z,א-ת]", text)
    return has_char


def check_text_for_illegal_labels(par):
    # capture if there is [start:start] or [end:end]
    error_pattern_st = (
        "("
        + defines.START_CHAR
        + "(?:(?!"
        + defines.END_CHAR
        + ").)*"
        + defines.START_CHAR
        + ")"
    )
    error_pattern_end = (
        "("
        + defines.END_CHAR
        + "(?:(?!"
        + defines.START_CHAR
        + ").)*"
        + defines.END_CHAR
        + ")"
    )
    error_pattern = "(" + error_pattern_st + "|" + error_pattern_end + ")"
    if re.search(error_pattern, par):
        print("\nstring gave ERROR\n{}".format(par))
        return 1
    return 0




def add_sent_column_for_labels():
    global sent_db
    sent_db["sent_idx_in_nar"] = (
        sent_db[sent_db["is_nar"] == 1].groupby(["doc_idx", "nar_idx"]).cumcount() + 1
    )
    sent_db["nar_len_in_sent"] = (
        sent_db[sent_db["is_nar"] == 1]
        .groupby(["doc_idx", "nar_idx"])["sent_idx_in_nar"]
        .transform("max")
    )

    sent_db["sent_idx_out_nar"] = (
        sent_db[sent_db["is_nar"] == 0].groupby(["doc_idx", "block_idx"]).cumcount() + 1
    )
    sent_db["fist_sent_in_nar"] = np.where(sent_db["sent_idx_in_nar"] == 1, True, False)
    sent_db["last_sent_in_nar"] = np.where(
        sent_db["sent_idx_in_nar"] == sent_db["nar_len_in_sent"], True, False
    )


def get_dummies_is_client():
    global sent_db
    sent_db["is_client"] = np.where(sent_db["par_type"] == "client", 1, 0)


def handle_short_sent_in_block(block):
    handled_block = block
    block_len = len(handled_block)
    curr_dot_idx = block_len - 1
    prev_dot_idx = 0
    # step = 0
    while curr_dot_idx < block_len:
        # step+=1
        curr_dot_idx = find_dot_idx(handled_block, prev_dot_idx + 1)
        # print("step {} curr_dot_idx {} prev_dot_idx {}".format(step,curr_dot_idx,prev_dot_idx))
        if curr_dot_idx < 1:
            break
        focus = handled_block[prev_dot_idx:curr_dot_idx]
        word_count = count_words(focus)
        # print ("\tword count in {} is {}".format(word_count,focus))
        if word_count > 0 and word_count <= defines.MIN_SENT_LEN:
            handled_block = replace_char_at_index(handled_block, curr_dot_idx)
        prev_dot_idx = curr_dot_idx
    return handled_block


def replace_char_at_index(org_str, index, replacement=" "):
    """Replace character at index in string org_str with the
    given replacement character."""
    new_str = org_str
    if index < len(org_str):
        new_str = org_str[0:index] + replacement + org_str[index + 1 :]
    return new_str


def find_dot_idx(text, start):
    return text.find(".", start)


def count_words(text):
    return len(text.split())

def remove_multi_x(text):
    return re.sub(r'X{3,4}.*?X{1,4}|X{3,}',' XXX ',text)


def split_block_to_sentences(text_, merge_short):
    text = text_
    text = remove_lr_annotation(text)
    # important to remove before we split into sentences
    text = replace_brackets(text)
    text = remove_multi_dots(text)
    text = remove_multi_x(text)
    text = remove_symbols(text)
    text = unify_numbers(text)
    text = replase_shekel_char(text)
    if merge_short:
        text = handle_short_sent_in_block(text)

    sent_list = tokenize.sent_tokenize(text)
    for i, item in enumerate(sent_list):
        clean_item = clean_text(item)
        check_text_for_symbols(clean_item)

        if (
            len(clean_item) != 0 and clean_item.isspace() == False
        ):  # disregard empty strings
            sent_list[i] = clean_item
    return sent_list

def replase_shekel_char(text):
    return re.sub(r'\u20AA',defines.SHEKEL,text)

def remove_multi_dots(text):
    text_ = re.sub(r"\A\.", "", text)  # replase ".נקודה בתחילת משפט"
    text_ = re.sub(r"\.+?\?", "?", text_)  # replace ?.. with ?
    # return re.sub(r'\.{2,3}', '',text_) # replace .. and ... with whitespace
    return re.sub(r"\.{2,3}", ".", text_)  # replace .. and ... with .

def unify_numbers(text,unify=' 123 '):
    return re.sub(r'[0-9]{1,}',unify,text)

def check_text_for_symbols(text):
    if re.search(r"[\t,\\t]", text):
        print("ERROR text has bad symbols {}".format(text))
        quit()


def add_sentences_of_blocks_to_db(block_db_idx,merge_short_sent):
    global sent_db
    block_line = block_db.iloc[block_db_idx]
    block = block_line["text"]
    sent_list = split_block_to_sentences(block,merge_short_sent)
    for i, sentence in enumerate(sent_list):
        if not text_contains_char(sentence):
            # print("Sentence wihtout char! \n {}".format(sentence))
            add_to_debug_df([("block_no_char", block), ("sent_no_char", sentence)])
            continue
        curr_db_idx = sent_db.shape[0]
        sent_db.loc[curr_db_idx, "is_question"] = 1 if "?" in sentence else 0
        sent_db.loc[curr_db_idx,'text'] = re.sub(r'\?','',sentence)
        sent_db.loc[curr_db_idx, "sent_idx_in_block"] = i
        sent_db.loc[curr_db_idx, "block_idx"] = block_db_idx
        sent_db.loc[curr_db_idx, "is_nar"] = block_line["is_nar"]
        sent_db.loc[curr_db_idx, "doc_idx"] = block_line["doc_idx"]
        sent_db.loc[curr_db_idx, "par_db_idx"] = block_line["par_db_idx"]
        sent_db.loc[curr_db_idx, "par_idx_in_doc"] = block_line["par_idx_in_doc"]
        sent_db.loc[curr_db_idx, "par_pos_in_doc"] = block_line["par_pos_in_doc"]
        sent_db.loc[curr_db_idx, "par_type"] = block_line["par_type"]
        sent_db.loc[curr_db_idx, "block_type"] = block_line["block_type"]
        sent_db.loc[curr_db_idx, "nar_idx"] = block_line["nar_idx"]
        sent_db.loc[curr_db_idx, "sent_len"] = len(sentence)


def add_to_debug_df(tupple_list):
    global debug_db
    idx = debug_db.shape[0]
    for tupple in tupple_list:
        debug_db.loc[idx, tupple[0]] = tupple[1]


def split_par_to_blocks_keep_order(par_db_idx):
    par = par_db.loc[par_db_idx, "text"]
    startNum = par.count(defines.START_CHAR)
    endNum = par.count(defines.END_CHAR)
    block_list = []  # holds tupple ("tag", "block string")
    tag = ""
    outside_nar = ""
    splited = []

    if startNum == 0 and endNum == 0:  # text is missing start and end symbols
        if par_db.loc[par_db_idx, "is_nar"] == 0:  # entire paragraph is not narrative
            tag = "not_nar"
        else:  # entire paragraph is narrative
            tag = "middle"
        block_list.insert(0, (tag, par))
    else:
        # used for keeping original order between blocks
        splited = re.split("&|#", par)
        splited_clean = splited.copy()
        for i, block in enumerate(splited):
            if block_has_summary(block):  # TBD handle story summary
                block,summ = extract_narrative_summary(block)
            splited_clean[i] = clean_text(block)
        my_regex = {
            # [start:end]
            "whole": defines.START_CHAR + ".*?" + defines.END_CHAR,
            "start": defines.START_CHAR + ".*",  # [start:]
            "end": ".*" + defines.END_CHAR,  # [:end]
        }
        outside_nar = par
        for tag, regex in my_regex.items():
            nar_blocks = re.findall(regex, outside_nar)
            for j, block in enumerate(nar_blocks):
                if len(block) != 0:
                    block_idx = get_index_of_block_in_par(
                        splited_clean, block, par_db_idx
                    )
                    # erase narrative blocks from splited paragraph
                    splited[block_idx] = ""
                    block_list.insert(block_idx, (tag, block))
            outside_nar = re.sub(regex, "", outside_nar)

        # handle the rest items in list - that must be non-narrative
        for i, block in enumerate(splited):
            if len(block) != 0:
                if block_has_summary(block):
                    block,summ = extract_narrative_summary(block)  # TBD handle story summary
                block_idx = get_index_of_block_in_par(splited_clean, block, par_db_idx)
                block_list.insert(block_idx, ("not_nar", block))
    check_block_list(par, splited, block_list)
    return block_list


def block_has_summary(text):
    if not "%" in text:
        return 0
    else:
        occur = text.count("%")
        if occur % 2 == 0:
            return 1
        else:
            add_to_debug_df([("odd_%", text)])
            return 0


def check_block_list(par, splited, block_list):
    global debug_db
    for i, block in enumerate(block_list):
        if not isinstance(block[1], str):
            spl = "".join(splited) if isinstance(splited, list) else "empty"
            add_to_debug_df(
                [
                    ("block_idx", i),
                    ("empty_block", str(block)),
                    ("splited", spl),
                    ("par", par),
                ]
            )
            # TBD remove
            # debug_db.to_csv(os.path.join(os.getcwd(),defines.PATH_TO_DFS,"debug_db.csv"),index=False)
            # sys.exit()


def get_index_of_block_in_par(splited, block, par_db_idx):
    cl_block = clean_text(block)
    if not cl_block in splited:
        print("{} \n par[{}]not in \n{}".format(par_db_idx, cl_block, splited))
        return -1
    else:
        return splited.index(cl_block)


def get_last_nar_idx_from_block_db():
    global block_db
    return block_db["nar_idx"].max()


def add_blocks_of_par_to_db(par_db_idx):
    global par_db, block_db
    is_nar = 0
    block_list = split_par_to_blocks_keep_order(par_db_idx)
    par_db_line = par_db.iloc[par_db_idx]
    for i, tupple in enumerate(block_list):
        curr_db_idx = block_db.shape[0]
        curr_nar_idx = 0 if curr_db_idx == 0 else get_last_nar_idx_from_block_db()
        if tupple[0] in ["start", "whole"]:
            curr_nar_idx += 1
        is_nar = 1 if tupple[0] != "not_nar" else 0
        block_db.loc[curr_db_idx, "text"] = tupple[1]
        block_db.loc[curr_db_idx, "is_nar"] = is_nar
        block_db.loc[curr_db_idx, "doc_idx"] = par_db_line["doc_idx"]
        block_db.loc[curr_db_idx, "par_idx_in_doc"] = par_db_line["par_idx_in_doc"]
        block_db.loc[curr_db_idx, "par_pos_in_doc"] = par_db_line["par_pos_in_doc"]
        block_db.loc[curr_db_idx, "par_db_idx"] = par_db_idx
        block_db.loc[curr_db_idx, "par_type"] = par_db_line["par_type"]
        block_db.loc[curr_db_idx, "block_type"] = tupple[0]
        block_db.loc[curr_db_idx, "nar_idx"] = curr_nar_idx if is_nar else 0


def read_csv(base_filename):
    return pd.read_csv("/".join([".", defines.PATH_TO_DFS, base_filename]))


def save_df_to_csv(df_name, is_global=True):
    path_to_save = "/".join([".", defines.PATH_TO_DFS, df_name])
    _df = globals().get(df_name, None)
    if _df is None:
        print("Can't find global df {}".format(df_name))
        return 1
    else:
        _df.to_csv("{}.csv".format(path_to_save), index=False)
        return 0


def save_doc_blocks(dir_name,doc_idx_from_name):
    global par_db, block_db
    block_db = pd.DataFrame()
    par_db = pd.read_csv(
        os.path.join(
            os.getcwd(),
            defines.PATH_TO_DFS,
            dir_name,
            "{:02d}_par_db.csv".format(doc_idx_from_name),
        )
    )
    for i in par_db.index:
        add_blocks_of_par_to_db(i)
    del par_db
    block_db.to_csv(
        os.path.join(
            os.getcwd(),
            defines.PATH_TO_DFS,
            dir_name,
            "{:02d}_block_db.csv".format(doc_idx_from_name),
        ),
        index=False,
    )
    del block_db
    # print("Doc {} blocks saved".format(doc_idx_from_name))



def calc_position_in_grp(x):
    return ((x+1)/len(x))


def save_doc_sentences(dir_name,doc_idx_from_name,merge_short_sent):
    global block_db, sent_db
    sent_db = pd.DataFrame()
    block_db = pd.read_csv(
        os.path.join(
            os.getcwd(),
            defines.PATH_TO_DFS,
            dir_name,
            "{:02d}_block_db.csv".format(doc_idx_from_name),
        )
    )
    for i in block_db.index:
        add_sentences_of_blocks_to_db(i,merge_short_sent)
    del block_db
    # add_sent_column_for_labels()
    get_dummies_is_client()
    sent_db["sent_idx_in_par"] = sent_db.groupby("par_idx_in_doc").cumcount()
    sent_db['sent_pos_in_par'] = sent_db.groupby('par_idx_in_doc')['sent_idx_in_par'].transform(calc_position_in_grp)
    sent_db['sent_pos_in_doc'] = (sent_db.index.values+1)/len(sent_db.index)
    sent_db.to_csv(
        os.path.join(
            os.getcwd(),
            defines.PATH_TO_DFS,
            dir_name,
            "{:02d}_sent_db.csv".format(doc_idx_from_name),
        ),
        index=False,
    )
    doc_db_update_stat(
        get_dbIdx_by_docIdx(doc_idx_from_name), "sent_count", len(sent_db.index)
    )
    doc_db_update_stat(
        get_dbIdx_by_docIdx(doc_idx_from_name),
        "nar_sent_count",
        len(sent_db[sent_db["is_nar"] == 1].index),
    )
    print("{} sentences".format(len(sent_db.index)), end = ' ')
    del sent_db


def save_doc_paragraphs(dir_name,doc_idx_from_name):
    global par_db, doc_db
    inside_narrative = 0
    doc_path = doc_db.loc[get_dbIdx_by_docIdx(doc_idx_from_name), "path"].values[0]
    if os.path.isfile(doc_path):
        doc = docx.Document(doc_path)
    else:
        print("Error: doc {} does not exist".format(doc_idx_from_name))
        return
    par_db = pd.DataFrame()
    for i, par in enumerate(doc.paragraphs):
        curr_par_db_idx = par_db.shape[0]
        text, par_type = get_par_type_erase(par.text)
        if not par_type in ["client", "therapist"]:
            print(
                "ERROR got par_type is {}, doc {}, par {} text\{}".format(
                    par_type, doc_idx_from_name, i, text
                )
            )
            os.exit()
        if len(text) == 0:
            continue
        par_db.loc[curr_par_db_idx, "doc_idx"] = doc_idx_from_name
        par_db.loc[curr_par_db_idx, "text"] = text
        par_db.loc[curr_par_db_idx, "par_len"] = len(text)
        par_db.loc[curr_par_db_idx, "par_type"] = par_type
        par_db.loc[curr_par_db_idx, "par_idx_in_doc"] = i
        if defines.START_CHAR in par.text:
            inside_narrative = 1
        par_db.loc[curr_par_db_idx, "is_nar"] = inside_narrative
        # if [...# ] or [ ...&...#]
        if par.text.rfind(defines.END_CHAR) > par.text.rfind(defines.START_CHAR):
            inside_narrative = 0
    par_db["par_pos_in_doc"] = (par_db.index.values+1)/len(par_db.index)
    par_db.to_csv(
        os.path.join(
            os.getcwd(),
            defines.PATH_TO_DFS,
            dir_name,
            "{:02d}_par_db.csv".format(doc_idx_from_name),
        ),
        index=False,
    )
    doc_db_update_stat(
        get_dbIdx_by_docIdx(doc_idx_from_name), "par_count", len(doc.paragraphs)
    )
    del par_db
    # print("Doc {} paragraphs saved".format(doc_idx_from_name))


def doc_db_update_stat(idx, val_name, value):
    global goc_db
    doc_db.loc[idx, val_name] = value


def save_all_docs_paragraphs():
    global doc_db
    for doc_idx in doc_db.index:
        save_doc_paragraphs(doc_idx)


def get_dbIdx_by_docIdx(doc_idx):
    global doc_db
    # print ("Doc idx {} db idx {}".format(doc_idx,doc_db[doc_db['doc_idx_from_name']==doc_idx].index.values))
    return doc_db[doc_db["doc_idx_from_name"] == doc_idx].index.values


def save_docs_db(doc_path_list=None,dir_name=""):
    global doc_db
    doc_db_path = os.path.join(os.getcwd(), defines.PATH_TO_DFS, dir_name, "doc_db.csv")
    if os.path.isfile(doc_db_path):
        os.remove(doc_db_path)
    print("Creating doc_db")
    doc_db = pd.DataFrame()
    if doc_path_list is None:
        doc_path_list = get_labeled_files()
    for path in doc_path_list:
        add_doc_to_db(path)
    doc_db["doc_idx_from_name"] = doc_db["doc_idx_from_name"].astype(int)
    doc_db.to_csv(doc_db_path, index=False)
    del doc_db


def remove_punctuation(_text):
    punct = re.sub("\?", "", string.punctuation)  # keep ?
    text = _text.translate(str.maketrans("", "", punct))
    return text


def get_labeled_files():
    doc_list = glob.glob(os.path.join(os.path.join(os.getcwd(), "./tmp/*")))
    doc_list.sort()
    return doc_list


def add_doc_to_db(path):
    global doc_db
    if not os.path.isfile(path):
        print("ERROR: file {} does not exists".format(path))
        return
    file_name = os.path.basename(path)
    doc_db_idx = doc_db.shape[0]
    doc_idx_from_name = common_utils.get_doc_idx_from_name(path)
    doc_db.loc[doc_db_idx, "path"] = path
    doc_db.loc[doc_db_idx, "file_name"] = file_name
    doc_db.loc[doc_db_idx, "doc_idx_from_name"] = int(doc_idx_from_name)


def replace_brackets(text, replace=""):  # remove [..] , (..), [..[.]..]
    return re.sub(r"\([^(]*?\)|\[[^[]*?\]", replace, text)


def clean_text(text):
    text_, _ = extract_narrative_summary(text)
    text_ = remove_punctuation(text_)
    return text_


def remove_symbols(text):
    return re.sub(r"[@#&<>\*\t\\t]", "", text)


def extract_narrative_summary(text):
    text_ = text
    summary = re.findall("%.*?%", text_)
    for i in summary:
        text_ = re.sub(i, "", text_)
    return text_, summary


def remove_lr_annotation(text):
    return re.sub(r"\(L[0-9].*?\-[A-Z]{1}\)", "", text)


def get_par_type_erase(par):
    par_type = "no_mark"
    if "CLIENT" in par[:20]:  # search for a tag in the begginning of a line
        par = re.sub("CLIENT(.*?:|)", "", par)
        par_type = "client"
    if "THERAPIST" in par[:20]:  # search for a tag in the begginning of a line
        par = re.sub("THERAPIST(.*?:|)", "", par)
        par_type = "therapist"
    check_text_for_illegal_labels(par)
    return par, par_type


def parse_all_docs(dir_name,merge_short_sent,doc_path_list=None):
    global doc_db, debug_db
    save_docs_db(doc_path_list,dir_name)
    doc_db_path = os.path.join(os.getcwd(), defines.PATH_TO_DFS, dir_name,"doc_db.csv")
    doc_db = pd.read_csv(doc_db_path)
    debug_db = pd.DataFrame()
    doc_indices = doc_db["doc_idx_from_name"].values
    doc_indices.sort()
    for i, doc_idx in enumerate(doc_indices):
        parse_doc(dir_name,int(doc_idx),merge_short_sent)
        print("{}".format(i),end=' ')
    doc_db.to_csv(
        os.path.join(os.getcwd(), defines.PATH_TO_DFS, dir_name,"doc_db.csv"), index=False
    )
    debug_db.to_csv(
        os.path.join(os.getcwd(), defines.PATH_TO_DFS, dir_name, "debug_db.csv"), index=False
    )
    del doc_db
    del debug_db


def parse_doc(dir_name,doc_idx, merge_short_sent, single=False):
    global doc_db, debug_db
    if single:
        doc_db_path = os.path.join(os.getcwd(), defines.PATH_TO_DFS, dir_name, "doc_db.csv")
        doc_db = pd.read_csv(doc_db_path)
        debug_db = pd.DataFrame()
    save_doc_paragraphs(dir_name,doc_idx)
    save_doc_blocks(dir_name,doc_idx)
    save_doc_sentences(dir_name,doc_idx,merge_short_sent)
    if single:
        debug_db.to_csv(
            os.path.join(os.getcwd(), defines.PATH_TO_DFS, dir_name, "debug_db.csv"), index=False
        )
        doc_db.to_csv(
            os.path.join(os.getcwd(), defines.PATH_TO_DFS, dir_name,"doc_db.csv"), index=False
        )


def add_new_doc(path,dir_name):
    global doc_db
    doc_db_path = os.path.join(os.getcwd(), defines.PATH_TO_DFS, dir_name, "doc_db.csv")
    doc_prefix = int(os.path.basename(path).split("_")[0])
    doc_db = pd.read_csv(doc_db_path)
    # if doc_prefix in doc_db['doc_idx_from_name'].values:
    #     print( "Doc with idx {} already parsed".format(doc_prefix))
    # else:
    #     add_doc_to_db(path)
    add_doc_to_db(path)
    doc_db.to_csv(doc_db_path, index=False)
